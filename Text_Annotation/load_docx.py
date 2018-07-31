import os
import docx
import jieba
from jieba import posseg
import re

jieba.setLogLevel('WARN')

DIR = os.path.dirname(os.path.abspath(__file__))

# 看着玩玩，实际只用了里面的数字
identifier = {1: 'n_S', 2: 'n_B', 3: 'n_M', 4: 'n_E',
              5: 'v_S', 6: 'v_B', 7: 'v_M', 8: 'v_E',
              9: 'unknow'}


def docx2data(file_path=None):
    '''

    :param file_path: 文件路径
    :return:
    '''
    document = docx.Document(file_path)

    texts, targets = [], []
    for i in range(len(document.paragraphs)):
        if document.paragraphs[i].text != '':
            text_line = document.paragraphs[i].text
            text_sentences=re.findall(pattern='[^，。；]+[，。；]',string=text_line)
            for text_sentence in text_sentences:
                target = []
                text_line_cut = posseg.lcut(text_sentence)
                for word_pair in text_line_cut:
                    word = word_pair.word
                    if word_pair.flag == 'n':
                        # 字标识为nS
                        if len(word) == 1:
                            target += [1]
                        # 词语标识为nB,nM,nE
                        else:
                            target += ([2] + [3] * (len(word) - 2) + [4])
                    elif word_pair.flag == 'v':
                        # 字标识为vS
                        if len(word) == 1:
                            target += [5]
                        # 词语标识为vB,vM,vE
                        else:
                            target += ([6] + [7] * (len(word) - 2) + [8])
                    else:
                        # 标注为unknow
                        target += ([9] * len(word))

                texts.append(text_sentence)
                targets.append(target)

    return texts, targets


def load_docx():
    texts, targets = [], []
    folder_names = ['材料清单', '法律法规', '医疗器械']
    for folder_name in folder_names:
        folder_path = DIR + '/data/%s' % (folder_name)
        file_paths = os.listdir(folder_path)
        for file_path in file_paths:
            texts_one, targets_one = docx2data(file_path=folder_path + '/' + file_path)
            texts += texts_one
            targets += targets_one

    return texts, targets


if __name__ == '__main__':
    texts, targets = load_docx()
